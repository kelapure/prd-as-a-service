from __future__ import annotations

import io
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from pathlib import Path

import fitz
from docx import Document


MAX_TOTAL_BYTES = 25 * 1024 * 1024
MAX_PAGES = 200
MAX_TEXT_CHARS = 250_000
MAX_IMAGES = 12
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}
IMAGE_MIME_BY_SUFFIX = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".gif": "image/gif",
    ".webp": "image/webp",
}


class InputError(ValueError):
    pass


@dataclass
class ExtractedImage:
    data: bytes
    media_type: str
    locator: str


@dataclass
class ExtractedDocument:
    name: str
    file_type: str
    text: str
    evidence_text: str = ""
    page_count: int | None = None
    sections: list[str] = field(default_factory=list)
    images: list[ExtractedImage] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def validate_filename(name: str) -> str:
    safe_name = Path(name).name
    extension = Path(safe_name).suffix.lower()
    if extension not in ALLOWED_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_EXTENSIONS))
        raise InputError(f"Unsupported file type. Use {allowed}; legacy .doc is not supported.")
    return safe_name


def extract_document(name: str, data: bytes) -> ExtractedDocument:
    safe_name = validate_filename(name)
    extension = Path(safe_name).suffix.lower()
    if not data:
        raise InputError(f"{safe_name} is empty")
    if extension == ".pdf":
        return _extract_pdf(safe_name, data)
    if extension == ".docx":
        return _extract_docx(safe_name, data)
    return _extract_text(safe_name, data, extension)


def extract_pasted_text(text: str) -> ExtractedDocument:
    normalized = text.strip()
    if len(normalized) < 50:
        raise InputError("Pasted PRD text must contain at least 50 characters")
    if len(normalized) > MAX_TEXT_CHARS:
        raise InputError(f"Pasted PRD text exceeds {MAX_TEXT_CHARS:,} characters")
    return ExtractedDocument(
        name="Pasted PRD",
        file_type="text/plain",
        text="[Source: Pasted PRD]\n" + normalized,
        evidence_text=normalized,
        sections=_headings(normalized),
    )


def _extract_text(name: str, data: bytes, extension: str) -> ExtractedDocument:
    try:
        text = data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise InputError(f"{name} must be UTF-8 text") from exc
    if len(text) > MAX_TEXT_CHARS:
        raise InputError(f"{name} exceeds {MAX_TEXT_CHARS:,} characters")
    if len(text.strip()) < 50:
        raise InputError(f"{name} does not contain enough PRD content")
    return ExtractedDocument(
        name=name,
        file_type="text/markdown" if extension == ".md" else "text/plain",
        text=f"[Source: {name}]\n{text.strip()}",
        evidence_text=text.strip(),
        sections=_headings(text),
    )


def _extract_pdf(name: str, data: bytes) -> ExtractedDocument:
    if not data.startswith(b"%PDF"):
        raise InputError(f"{name} does not contain a valid PDF signature")
    try:
        document = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise InputError(f"{name} could not be opened as a PDF") from exc
    if document.page_count > MAX_PAGES:
        raise InputError(f"{name} has {document.page_count} pages; the limit is {MAX_PAGES}")

    pages: list[str] = []
    evidence_pages: list[str] = []
    images: list[ExtractedImage] = []
    warnings: list[str] = []
    for index, page in enumerate(document):
        page_number = index + 1
        text = page.get_text("text").strip()
        pages.append(f"[Page {page_number}]\n{text or '[No extractable text on this page]'}")
        if text:
            evidence_pages.append(text)
        should_render = len(text) < 80 or bool(page.get_images(full=True))
        if should_render and len(images) < MAX_IMAGES:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.35, 1.35), alpha=False)
            images.append(
                ExtractedImage(
                    data=pixmap.tobytes("png"),
                    media_type="image/png",
                    locator=f"{name}, page {page_number}",
                )
            )

    combined = f"[Source: {name}]\n" + "\n\n".join(pages)
    if len(combined) > MAX_TEXT_CHARS:
        raise InputError(f"Extracted text from {name} exceeds {MAX_TEXT_CHARS:,} characters")
    if not any(page.strip() and "[No extractable" not in page for page in pages) and not images:
        raise InputError(f"{name} has no readable text or renderable pages")
    if len(images) == MAX_IMAGES:
        warnings.append("Figure-aware review was limited to the first 12 relevant page images.")
    return ExtractedDocument(
        name=name,
        file_type="application/pdf",
        text=combined,
        evidence_text="\n\n".join(evidence_pages),
        page_count=document.page_count,
        sections=_headings(combined),
        images=images,
        warnings=warnings,
    )


def _extract_docx(name: str, data: bytes) -> ExtractedDocument:
    if not data.startswith(b"PK"):
        raise InputError(f"{name} does not contain a valid DOCX signature")
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
        if "word/document.xml" not in archive.namelist():
            raise InputError(f"{name} is not a valid DOCX document")
        document = Document(io.BytesIO(data))
    except (zipfile.BadZipFile, ValueError) as exc:
        raise InputError(f"{name} could not be opened as DOCX") from exc

    declared_pages: int | None = None
    if "docProps/app.xml" in archive.namelist():
        try:
            app_properties = ET.fromstring(archive.read("docProps/app.xml"))
            pages_element = next(
                (element for element in app_properties.iter() if element.tag.endswith("}Pages") or element.tag == "Pages"),
                None,
            )
            if pages_element is not None and pages_element.text:
                declared_pages = int(pages_element.text)
                if declared_pages <= 0:
                    declared_pages = None
        except (ET.ParseError, ValueError):
            declared_pages = None
    if declared_pages is not None and declared_pages > MAX_PAGES:
        raise InputError(f"{name} has {declared_pages} pages; the limit is {MAX_PAGES}")

    blocks: list[str] = [f"[Source: {name}]"]
    evidence_blocks: list[str] = []
    sections: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = paragraph.style.name.lower() if paragraph.style else ""
        if style.startswith("heading") or style == "title":
            sections.append(text)
            blocks.append(f"\n[Section: {text}]")
            evidence_blocks.append(text)
        else:
            blocks.append(text)
            evidence_blocks.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        blocks.append(f"\n[Table {table_index}]")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            blocks.append(row_text)
            evidence_blocks.extend(cell.text.strip() for cell in row.cells if cell.text.strip())

    images: list[ExtractedImage] = []
    warnings: list[str] = []
    if declared_pages is None:
        warnings.append("DOCX page count was unavailable; byte and extracted-text limits were enforced.")
    for member in sorted(archive.namelist()):
        if not member.startswith("word/media/"):
            continue
        suffix = Path(member).suffix.lower()
        media_type = IMAGE_MIME_BY_SUFFIX.get(suffix)
        if not media_type:
            warnings.append(f"Skipped unsupported embedded media format {suffix or 'unknown'}.")
            continue
        if len(images) >= MAX_IMAGES:
            warnings.append("Figure-aware review was limited to the first 12 embedded images.")
            break
        images.append(
            ExtractedImage(
                data=archive.read(member),
                media_type=media_type,
                locator=f"{name}, embedded figure {len(images) + 1}",
            )
        )

    combined = "\n".join(blocks).strip()
    if len(combined) > MAX_TEXT_CHARS:
        raise InputError(f"Extracted text from {name} exceeds {MAX_TEXT_CHARS:,} characters")
    if len(combined) < 50 and not images:
        raise InputError(f"{name} does not contain enough readable PRD content")
    return ExtractedDocument(
        name=name,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        text=combined,
        evidence_text="\n".join(evidence_blocks),
        page_count=declared_pages,
        sections=sections,
        images=images,
        warnings=warnings,
    )


def _headings(text: str) -> list[str]:
    headings: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            value = stripped.lstrip("#").strip()
            if value:
                headings.append(value[:160])
        elif stripped.startswith("[Section:") and stripped.endswith("]"):
            headings.append(stripped[9:-1].strip()[:160])
    return headings[:100]
